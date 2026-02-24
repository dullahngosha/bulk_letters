import os
import zipfile
import pandas as pd
from flask import Flask, render_template, request, send_file, redirect, url_for
from docx import Document
from io import BytesIO

app = Flask(__name__, template_folder='templates', static_folder='static')

def replace_placeholders(doc, data_row):
    """Inabadilisha tagi [TAGI] kwenda thamani halisi."""
    for paragraph in doc.paragraphs:
        for key, value in data_row.items():
            tag = f"[{key}]"
            if tag in paragraph.text:
                paragraph.text = paragraph.text.replace(tag, str(value))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data_row.items():
                        tag = f"[{key}]"
                        if tag in paragraph.text:
                            paragraph.text = paragraph.text.replace(tag, str(value))

@app.route('/')
def index():
    return render_template('index.html')

# Hapa ndipo tulirekebisha: Tumeruhusu GET na POST
@app.route('/process', methods=['GET', 'POST'])
def process_files():
    # Ikiwa mtu ameingia URL hii moja kwa moja bila kupakia faili
    if request.method == 'GET':
        return redirect(url_for('index'))

    try:
        word_file = request.files.get('word_template')
        excel_file = request.files.get('excel_data')
        filename_col = request.form.get('filename_column', '').strip()

        if not word_file or not excel_file:
            return "Tafadhali pakia faili zote mbili kabla ya kuendelea.", 400

        # Soma Excel
        try:
            df = pd.read_excel(excel_file)
        except Exception:
            return "Faili la Excel halisomeki. Hakikisha ni .xlsx safi.", 400

        if filename_col not in df.columns:
            return f"Kosa: Safu (Column) yenye jina '{filename_col}' haiko kwenye Excel yako.", 400

        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED) as zip_file:
            for index, row in df.iterrows():
                word_file.seek(0)
                doc = Document(word_file)
                replace_placeholders(doc, row.to_dict())
                
                doc_io = BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                # Safisha jina la faili
                raw_name = str(row[filename_col])
                safe_name = "".join([c for c in raw_name if c.isalnum() or c in (' ', '-', '_')]).rstrip()
                individual_filename = f"Barua_{safe_name}.docx"
                zip_file.writestr(individual_filename, doc_io.getvalue())

        zip_buffer.seek(0)
        return send_file(
            zip_buffer, 
            mimetype='application/zip', 
            as_attachment=True, 
            download_name='Ngosha_Bulk_Letters.zip'
        )

    except Exception as e:
        return f"Kuna tatizo limetokea: {str(e)}", 500

if __name__ == '__main__':
    app.run(debug=True)
