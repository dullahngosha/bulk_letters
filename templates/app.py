import os
import zipfile
import pandas as pd
from flask import Flask, render_template, request, send_file
from docx import Document
from io import BytesIO

# Tunahakikisha Flask inajua folder la templates lilipo
app = Flask(__name__, template_folder='templates')

def replace_placeholders(doc, data_row):
    """Inabadilisha tagi kama [JINA] kwenda kwenye thamani halisi."""
    for paragraph in doc.paragraphs:
        for key, value in data_row.items():
            tag = f"[{key}]"
            if tag in paragraph.text:
                paragraph.text = paragraph.text.replace(tag, str(value))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data_row.items():
                        tag = f"[{key}]"
                        if tag in paragraph.text:
                            paragraph.text = paragraph.text.replace(tag, str(value))

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process_files():
    try:
        word_file = request.files['word_template']
        excel_file = request.files['excel_data']
        filename_col = request.form.get('filename_column', 'JINA')

        # Soma Excel
        df = pd.read_excel(excel_file)
        
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED) as zip_file:
            for index, row in df.iterrows():
                doc = Document(word_file)
                replace_placeholders(doc, row.to_dict())
                
                doc_io = BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                # Safisha jina la faili
                name = str(row.get(filename_col, index)).replace(" ", "_")
                individual_filename = f"Barua_{name}.docx"
                zip_file.writestr(individual_filename, doc_io.getvalue())

        zip_buffer.seek(0)
        return send_file(
            zip_buffer, 
            mimetype='application/zip', 
            as_attachment=True, 
            download_name='Ngosha_Bulk_Letters.zip'
        )
    except Exception as e:
        return f"Kuna tatizo limetokea: {str(e)}"

if __name__ == '__main__':
    app.run(debug=True)